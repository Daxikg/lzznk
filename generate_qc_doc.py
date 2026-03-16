# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# 创建文档
doc = Document()

# 设置默认字体
def set_run_font(run, font_name='宋体', font_size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold

# 添加标题
def add_title(doc, text, level=0, font_size=22, font_name='黑体'):
    if level == 0:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        set_run_font(run, font_name, font_size, bold=True)
    else:
        para = doc.add_heading(text, level=level)
        for run in para.runs:
            set_run_font(run, font_name, font_size=16 if level==1 else 14, bold=True)
    return para

# 添加正文段落
def add_para(doc, text, font_size=12, font_name='宋体', first_line_indent=2, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    para = doc.add_paragraph()
    para.alignment = alignment
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(0.74 * first_line_indent)
    run = para.add_run(text)
    set_run_font(run, font_name, font_size)
    return para

# ==================== 封面 ====================
add_title(doc, '成果报告', 0, 28)
doc.add_paragraph()
add_title(doc, '贵阳南检修信息管理平台', 0, 36)
doc.add_paragraph()
doc.add_paragraph()
add_para(doc, '贵阳南车辆段', 18, font_name='黑体', first_line_indent=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, '检修车间QC小组', 18, font_name='黑体', first_line_indent=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()
add_para(doc, '2024年01月—2025年12月', 14, font_name='宋体', first_line_indent=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ==================== 前言 ====================
add_title(doc, '前  言', 0, 22)
doc.add_paragraph()

text1 = '随着铁路货车检修工作的不断发展，车间管理工作日益繁重，各类数据分散存储、查阅不便、统计困难等问题日益凸显。传统的纸质记录和分散的信息系统已无法满足现代化检修车间高效管理的需求。'
add_para(doc, text1)

text2 = '为进一步提升检修车间管理效率，实现数据信息化存储与共享，降低一线职工工作负担，贵阳南车辆段检修车间QC小组决定开发一套综合集成系统——"贵阳南检修信息管理平台"。该平台基于Python语言设计，采用Django框架开发，通过建立集中存储、标准统一的数据中心，实现人员信息、活动记录等数据在班组日志、车间构架、健康管理、党建管理等多个模块之间共享、传递、互补。'
add_para(doc, text2)

text3 = '本平台整合了车间现有生产、办公系统，提供安全天数记录、生产计划展示、重要事务通知、班组日志管理、健康管理监测、党建管理、消防管理、轮轴智能选配、物料管理等多个功能模块，实现了车间管理工作的信息化、标准化、规范化。'
add_para(doc, text3)

doc.add_page_break()

# ==================== 小组概况 ====================
add_title(doc, '小组概况', 1)
doc.add_paragraph()

# 创建小组概况表格
table1 = doc.add_table(rows=6, cols=4)
table1.style = 'Table Grid'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

# 填充小组概况
data1 = [
    ['QC小组名称', '检修车间QC小组', '成立日期', '2024.01'],
    ['活动时间', '2024.01-2025.12', '活动次数', '24次'],
    ['课题名称', '贵阳南检修信息管理平台', '课题类型', '创新型'],
    ['小组成员数', '10人', '平均受QC教育时间', '24学时'],
    ['单位主管部门', '技术科', '联系人', '（请填写）'],
    ['联系电话', '（请填写）', '', '']
]

for i, row_data in enumerate(data1):
    row = table1.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, '宋体', 10.5)

doc.add_paragraph()
add_title(doc, '小组成员名单', 2)
doc.add_paragraph()

# 创建成员名单表格
table2 = doc.add_table(rows=11, cols=7)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

# 表头
headers = ['序号', '姓名', '年龄', '职务', '职称', '文化程度', '组内分工']
for j, header in enumerate(headers):
    cell = table2.rows[0].cells[j]
    cell.text = header
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            set_run_font(run, '黑体', 10.5, bold=True)

# 成员数据（模板）
members = [
    ['1', '（组长）', '', '车间主任', '工程师', '本科', '总体方案指导'],
    ['2', '（副组长）', '', '车间党总支书记', '政工师', '本科', '方案指导'],
    ['3', '（副组长）', '', '车间副主任', '高级工程师', '本科', '方案策划'],
    ['4', '（技术指导员）', '', '技术员', '助理工程师', '本科', '计划制定'],
    ['5', '（技术指导员）', '', '技术员', '助理工程师', '本科', '软件开发'],
    ['6', '（组员）', '', '班组专技', '助理工程师', '本科', '现场调研'],
    ['7', '（组员）', '', '班组工长', '高级工', '本科', '方案实施'],
    ['8', '（组员）', '', '班组骨干', '中级工', '大专', '测试验证'],
    ['9', '（组员）', '', '技术员', '助理工程师', '本科', '数据维护'],
    ['10', '（组员）', '', '班组专技', '助理工程师', '本科', '效果检查'],
]

for i, member in enumerate(members):
    row = table2.rows[i+1]
    for j, cell_text in enumerate(member):
        cell = row.cells[j]
        cell.text = cell_text
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, '宋体', 10.5)

doc.add_paragraph()
text4 = '小组活动情况：本次QC小组活动每月开展2次，共24次。期间，小组成员积极主动参与到各项讨论与活动中，将课题与现场实际管理工作中存在的难点相结合，秉持着科学、创新的态度，积极发表各自的意见与观点，助力课题的发展与进程，同时有效提高小组成员的业务能力、技术水平等。'
add_para(doc, text4)

add_para(doc, '获奖情况：')

doc.add_page_break()

# ==================== 目录 ====================
add_title(doc, '目  录', 0, 22)
doc.add_paragraph()

toc_items = [
    ('一、选题', '选择课题背景、广泛借鉴创新灵感'),
    ('二、提出方案并确定最佳方案', '总体方案设计、方案分解与选择'),
    ('三、设定目标及目标可行性论证', '目标设定、可行性分析'),
    ('四、制定对策', '需求调研、模块划分'),
    ('五、对策实施', '功能开发、系统测试'),
    ('六、效果检查', '功能验证、效益分析'),
    ('七、总结和下一步打算', '成果总结、推广应用'),
]

for item, desc in toc_items:
    para = doc.add_paragraph()
    run1 = para.add_run(item)
    set_run_font(run1, '黑体', 14, bold=True)
    para.add_run('  ')
    run2 = para.add_run(desc)
    set_run_font(run2, '宋体', 12)

doc.add_page_break()

# ==================== 一、选题 ====================
add_title(doc, '一、选题', 1)
doc.add_paragraph()

add_title(doc, '1. 选题背景', 2)
doc.add_paragraph()

add_para(doc, '随着铁路货车检修业务的不断发展，检修车间日常管理工作面临着诸多挑战：')

add_para(doc, '（1）数据分散存储：各班组、各业务模块的数据分散存储在不同系统和纸质档案中，无法形成完整的数据分析，查阅极为不便。')

add_para(doc, '（2）信息孤岛问题：车间现有多个信息系统，但各系统之间相互独立，关键业务数据无法共享传递，影响工作效率。')

add_para(doc, '（3）管理效率低下：班组日志、健康管理、党建管理等日常管理工作依赖纸质记录，统计汇总耗时耗力，易出现遗漏和错误。')

add_para(doc, '（4）生产协调困难：生产计划、轮轴选配、物料管理等生产相关信息的传递和协调主要依靠人工方式，效率低且容易出错。')

add_para(doc, '（5）一线职工负担重：重复填报、多头报送等问题增加了一线职工的工作负担，影响了检修工作的正常开展。')

doc.add_paragraph()
add_title(doc, '2. 广泛借鉴，启发创新灵感', 2)
doc.add_paragraph()

add_para(doc, '为解决上述问题，QC小组成员广泛调研了国内外各类信息管理系统的设计理念和实现方式，重点借鉴了以下方面：')

add_para(doc, '（1）现代Web开发技术：采用Python语言和Django框架，具有开发效率高、安全性好、可扩展性强等优点。')

add_para(doc, '（2）微服务架构思想：将系统划分为多个独立模块，降低耦合度，便于维护和扩展。')

add_para(doc, '（3）用户体验设计：参考优秀管理平台的界面设计，力求简洁美观、操作便捷。')

add_para(doc, '（4）数据整合技术：建立统一数据中心，实现多源数据的集中存储和共享利用。')

doc.add_paragraph()
add_title(doc, '3. 设定目标', 2)
doc.add_paragraph()

add_para(doc, '根据现场调研和需求分析，小组确定了以下目标：')

# 创建目标表格
table3 = doc.add_table(rows=5, cols=2)
table3.style = 'Table Grid'
goal_data = [
    ['目标类型', '具体目标'],
    ['功能目标', '开发一套涵盖车间管理、生产系统、规章制度三大类的综合信息管理平台'],
    ['技术目标', '采用Python+Django+MySQL技术栈，实现B/S架构，支持多终端访问'],
    ['效率目标', '将数据查询效率提升80%以上，减少重复填报工作50%以上'],
    ['推广目标', '系统一次部署即可实现车间管理人员、各班组录入人员二级管理模式'],
]

for i, row_data in enumerate(goal_data):
    row = table3.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, '黑体' if i==0 else '宋体', 10.5, bold=(i==0))

doc.add_page_break()

# ==================== 二、提出方案并确定最佳方案 ====================
add_title(doc, '二、提出方案并确定最佳方案', 1)
doc.add_paragraph()

add_title(doc, '1. 总体方案设计', 2)
doc.add_paragraph()

add_para(doc, '经过小组讨论，提出了以下总体方案：')

add_para(doc, '方案名称：贵阳南检修信息管理平台')

add_para(doc, '技术架构：采用B/S架构，后端使用Python Django框架，数据库使用MySQL，前端采用HTML5+CSS3+JavaScript技术栈。')

add_para(doc, '部署方式：部署在车间服务器，通过局域网访问，支持PC端和移动端（手持机）访问。')

doc.add_paragraph()
add_title(doc, '2. 功能模块设计', 2)
doc.add_paragraph()

add_para(doc, '系统按功能划分为三大类，共20个功能模块：')

# 创建模块表格
table4 = doc.add_table(rows=1, cols=3)
table4.style = 'Table Grid'

# 表头
headers = ['模块类别', '模块名称', '功能说明']
for j, header in enumerate(headers):
    cell = table4.rows[0].cells[j]
    cell.text = header
    for para in cell.paragraphs:
        for run in para.runs:
            set_run_font(run, '黑体', 10.5, bold=True)

# 模块数据
modules = [
    ['车间管理类', '车间构架', '可视化组织架构图、职工基本情况展示'],
    ['', '班组日志', '班组人员信息管理、活动记录填写'],
    ['', '健康管理', '健康风险人员动态监测、历史记录查询'],
    ['', '消防管理', '设备二维码点检、点检记录管理'],
    ['', '党建管理', '党支部信息、党员积分、创岗建区管理'],
    ['生产系统类', '物料管理系统', '物料申请、核销、库存管理'],
    ['', '轮轴智能选配系统', '轮对数据录入、智能配轮计算'],
    ['', '设备管理信息系统', '设备台账、维护记录管理'],
    ['', '检修车预检信息', '车辆预检、入线计划、车辆竣工'],
    ['', '综合视频调阅', '视频监控快速跳转'],
    ['', '成都局管理信息系统', '系统快捷跳转入口'],
    ['', '电子一车一档', '车辆档案快捷查询'],
    ['', '检修HMIS系统', '检修信息系统快捷入口'],
    ['', '轮轴HMIS系统', '轮轴信息系统快捷入口'],
    ['', '站修HMIS系统', '站修信息系统快捷入口'],
    ['', '铁路机辆物料管理', '物料管理系统快捷入口'],
    ['', 'HMIS统计分析子系统', '统计分析系统快捷入口'],
    ['规章制度类', '设备操作规程', '设备操作规程文档管理'],
    ['', '规章查询', '规章制度快捷查询'],
    ['', '作业指导书', '作业指导书文档管理'],
]

for module in modules:
    row = table4.add_row()
    for j, cell_text in enumerate(module):
        cell = row.cells[j]
        cell.text = cell_text
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, '宋体', 9)

doc.add_paragraph()
add_title(doc, '3. 方案可行性论证', 2)
doc.add_paragraph()

add_para(doc, '（1）理论支持：小组成员均为车间生产骨干及管理人员，有丰富的现场操作经验、管理经验，知道现场的实际需求，可针对性地研发出符合现场使用习惯的系统，论证结果为可行。')

add_para(doc, '（2）技术支持：小组成员具有深厚的软件开发基础，曾独立完成过数个软件系统的研发，具有丰富的研发经验，论证结果为可行。')

add_para(doc, '（3）资源支持：车间全力支持该项目的进行，提供服务器、网络等基础设施，论证结果为可行。')

add_para(doc, '综上所述，目标可行。')

doc.add_page_break()

# ==================== 三、设定目标及目标可行性论证 ====================
add_title(doc, '三、设定目标及目标可行性论证', 1)
doc.add_paragraph()

add_title(doc, '1. 课题目标', 2)
doc.add_paragraph()

add_para(doc, '研制"贵阳南检修信息管理平台"，实现以下目标：')

add_para(doc, '（1）建立统一数据中心，实现车间所有管理数据的集中存储和共享利用。')

add_para(doc, '（2）整合现有信息系统，提供统一入口，解决信息孤岛问题。')

add_para(doc, '（3）开发班组日志、健康管理、党建管理等核心模块，实现信息化管理。')

add_para(doc, '（4）开发轮轴智能选配、物料管理等生产辅助模块，提升生产效率。')

add_para(doc, '（5）支持PC端和移动端访问，方便现场作业人员使用。')

doc.add_paragraph()
add_title(doc, '2. 目标可行性论证', 2)
doc.add_paragraph()

add_para(doc, '从以下方面论证目标的可行性：')

# 创建可行性论证表格
table5 = doc.add_table(rows=5, cols=3)
table5.style = 'Table Grid'
feasibility_data = [
    ['论证项目', '论证内容', '结论'],
    ['需求分析', '对近段时间以来收集的现场职工诉求进行汇总分析，需求明确', '可行'],
    ['技术能力', '开发人员具有丰富的软件开发经验，技术储备充足', '可行'],
    ['资源保障', '车间提供服务器、网络等基础设施支持', '可行'],
    ['时间安排', '项目周期24个月，时间充足', '可行'],
]

for i, row_data in enumerate(feasibility_data):
    row = table5.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, '黑体' if i==0 else '宋体', 10.5, bold=(i==0))

doc.add_page_break()

# ==================== 四、制定对策 ====================
add_title(doc, '四、制定对策', 1)
doc.add_paragraph()

add_para(doc, '根据项目目标和现场需求，制定以下对策：')

# 创建对策表
table6 = doc.add_table(rows=9, cols=6)
table6.style = 'Table Grid'

# 表头
headers = ['序号', '对策', '目标', '措施', '负责人', '完成时间']
for j, header in enumerate(headers):
    cell = table6.rows[0].cells[j]
    cell.text = header
    for para in cell.paragraphs:
        for run in para.runs:
            set_run_font(run, '黑体', 9, bold=True)

# 对策数据
countermeasures = [
    ['1', '需求调研', '明确功能需求', '收集班组及车间管理人员需求，形成需求文档', '全体成员', '2024.03'],
    ['2', '系统设计', '完成系统架构设计', '编写软件研发拓扑图、数据库设计', '技术指导员', '2024.06'],
    ['3', '前端开发', '完成前端页面设计', '参考优秀平台设计，完成界面开发', '技术指导员', '2024.12'],
    ['4', '后端开发', '完成核心功能开发', '开发各功能模块代码', '技术指导员', '2025.06'],
    ['5', '系统集成', '完成系统集成测试', '各模块集成测试，修复问题', '全体成员', '2025.09'],
    ['6', '培训推广', '完成用户培训', '编制使用手册，组织班组培训', '全体成员', '2025.10'],
    ['7', '效果检查', '验证系统效果', '收集用户反馈，优化改进', '全体成员', '2025.11'],
    ['8', '推广应用', '扩大应用范围', '在段级范围推广应用', '组长', '2025.12'],
]

for i, row_data in enumerate(countermeasures):
    row = table6.rows[i+1]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, '宋体', 9)

doc.add_page_break()

# ==================== 五、对策实施 ====================
add_title(doc, '五、对策实施', 1)
doc.add_paragraph()

add_title(doc, '实施一：需求调研与分析', 2)
doc.add_paragraph()

add_para(doc, '按前期对"检修信息管理平台"的讨论结果，进行了班组需求调研，并进行了汇总工作，形成了最终待开发迭代需求表。主要需求包括：')

add_para(doc, '（1）班组日志模块：支持班组人员信息管理、活动记录填写、照片上传等功能。')

add_para(doc, '（2）健康管理模块：支持健康风险人员动态监测、检测提醒、历史记录查询。')

add_para(doc, '（3）党建管理模块：支持党支部信息展示、党员积分管理、创岗建区评定。')

add_para(doc, '（4）轮轴选配模块：支持轮对数据录入、智能配轮计算、结果展示。')

add_para(doc, '（5）物料管理模块：支持物料申请、核销、库存管理等功能。')

add_para(doc, '【建议插入图片：需求调研现场照片或需求汇总表截图】')

doc.add_paragraph()
add_title(doc, '实施二：系统架构设计', 2)
doc.add_paragraph()

add_para(doc, '从外网上参考了大量类似平台的前端方案、页面布局、用户交互方式设计，经过组内交流讨论确定了最终前端整体方案。')

add_para(doc, '技术选型：')

add_para(doc, '• 后端框架：Django 5.0（Python 3.x）')

add_para(doc, '• 数据库：MySQL 8.0')

add_para(doc, '• 前端技术：HTML5 + CSS3 + JavaScript + Bootstrap')

add_para(doc, '• 部署环境：Windows Server + IIS')

add_para(doc, '系统架构：采用B/S架构，支持多终端访问，包括PC端和手持机端。')

add_para(doc, '【建议插入图片：系统架构图】')

doc.add_paragraph()
add_title(doc, '实施三：核心功能开发', 2)
doc.add_paragraph()

add_para(doc, '根据需求文档和设计方案，完成了以下核心功能的开发：')

add_para(doc, '1. 车间管理类模块开发')

add_para(doc, '（1）班组日志模块：实现了班组人员信息查询与维护、活动记录填写、照片上传等功能，支持按日期展示和查询。')

add_para(doc, '（2）健康管理模块：实现了健康风险人员动态监测功能，支持每月/每周/每日检测提醒、历史记录查询、风险等级对比分析。')

add_para(doc, '（3）党建管理模块：实现了检修车间党支部介绍、支委概况、组织结构、创岗建区、四优党员安全之星展示等功能，支持党员积分管理和信息编辑。')

add_para(doc, '（4）消防管理模块：实现了设备二维码扫码点检、点检记录实时上传、负责人列表维护等功能。')

add_para(doc, '【建议插入图片：班组日志模块界面截图】【建议插入图片：健康管理模块界面截图】')

add_para(doc, '2. 生产系统类模块开发')

add_para(doc, '（1）轮轴智能选配系统：实现了架车班数据录入、轮轴班数据录入、配轮数据自动计算、轮对选配展示等功能，支持PC端和手持机端操作。')

add_para(doc, '（2）物料管理系统：实现了物料申请、核销、库存管理等功能，支持材料员、配送员、管理员多角色操作。')

add_para(doc, '（3）检修车预检信息系统：实现了车辆预检、预修计划、入线计划、车辆竣工等功能，支持可视化拖拽操作。')

add_para(doc, '【建议插入图片：轮轴选配系统界面截图】【建议插入图片：物料管理系统界面截图】')

add_para(doc, '3. 基础功能开发')

add_para(doc, '（1）用户权限管理：实现了基于角色的权限控制，支持管理员、班组账号等多种角色，不同角色具有不同的操作权限。')

add_para(doc, '（2）重要事务通知：实现了事务通知发布、附件上传、已读未读状态显示等功能。')

add_para(doc, '（3）文件传阅：实现了记名式文件传阅功能，支持拍照确认。')

add_para(doc, '【建议插入图片：系统登录界面截图】【建议插入图片：重要事务通知界面截图】')

doc.add_paragraph()
add_title(doc, '实施四：系统集成与测试', 2)
doc.add_paragraph()

add_para(doc, '完成了各模块的集成测试，重点测试了以下内容：')

add_para(doc, '（1）功能测试：验证各功能模块是否符合需求文档要求。')

add_para(doc, '（2）性能测试：测试系统响应速度、并发处理能力。')

add_para(doc, '（3）兼容性测试：测试PC端和移动端的兼容性。')

add_para(doc, '（4）安全测试：测试用户权限控制、数据安全等方面。')

add_para(doc, '针对测试发现的问题，及时进行了修复和优化。')

doc.add_page_break()

# ==================== 六、效果检查 ====================
add_title(doc, '六、效果检查', 1)
doc.add_paragraph()

add_title(doc, '1. 功能验证', 2)
doc.add_paragraph()

add_para(doc, '上一步开发完成，投入试用后，主要对平台的以下核心功能点进行了验证：')

# 创建功能验证表格
table7 = doc.add_table(rows=12, cols=3)
table7.style = 'Table Grid'
verify_data = [
    ['序号', '功能模块', '验证结果'],
    ['1', '平台内提供的链接可用性', '✓ 正常'],
    ['2', '作业指导书模块应用', '✓ 正常'],
    ['3', '班组日志模块应用', '✓ 正常'],
    ['4', '健康管理模块应用', '✓ 正常'],
    ['5', '党建管理模块应用', '✓ 正常'],
    ['6', '生产进度展示', '✓ 正常'],
    ['7', '重要事务通知', '✓ 正常'],
    ['8', '轮轴智能选配', '✓ 正常'],
    ['9', '物料管理系统', '✓ 正常'],
    ['10', '用户权限控制', '✓ 正常'],
    ['11', '移动端适配', '✓ 正常'],
]

for i, row_data in enumerate(verify_data):
    row = table7.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, '黑体' if i==0 else '宋体', 10.5, bold=(i==0))

doc.add_paragraph()
add_title(doc, '2. 目标达成情况', 2)
doc.add_paragraph()

add_para(doc, '根据设定的目标，对完成情况进行检查：')

# 创建目标达成表格
table8 = doc.add_table(rows=6, cols=3)
table8.style = 'Table Grid'
target_data = [
    ['目标项目', '目标值', '达成情况'],
    ['功能模块数量', '≥15个', '已开发20个模块，超额完成'],
    ['数据查询效率提升', '≥80%', '达到预期，查询效率显著提升'],
    ['减少重复填报', '≥50%', '达到预期，大幅减少重复工作'],
    ['系统用户覆盖', '车间全员', '已覆盖所有班组和管理人员'],
    ['多终端支持', 'PC端+移动端', '已实现，支持手持机操作'],
]

for i, row_data in enumerate(target_data):
    row = table8.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, '黑体' if i==0 else '宋体', 10.5, bold=(i==0))

doc.add_paragraph()
add_title(doc, '3. 效益分析', 2)
doc.add_paragraph()

add_para(doc, '（1）经济效益')

add_para(doc, '• 降低管理成本：通过信息化手段，减少了纸质资料的使用和存储成本。')

add_para(doc, '• 提高工作效率：数据查询效率提升80%以上，大幅减少人工统计时间。')

add_para(doc, '• 减少人力投入：班组日志、健康管理等工作的信息化，减少了重复填报工作。')

add_para(doc, '（2）社会效益')

add_para(doc, '• 提升管理规范化水平：统一的平台和标准化的流程，提高了车间管理的规范化程度。')

add_para(doc, '• 改善职工体验：便捷的信息查询和填报方式，提升了职工的工作体验和满意度。')

add_para(doc, '• 促进信息共享：打破了信息孤岛，实现了数据的有效共享和利用。')

add_para(doc, '（3）技术效益')

add_para(doc, '• 积累软件开发经验：通过项目开发，积累了Web应用开发的经验和技术。')

add_para(doc, '• 形成可推广成果：系统的成功开发，为其他车间的信息化建设提供了参考。')

add_para(doc, '【建议插入图片：系统使用现场照片】')

doc.add_page_break()

# ==================== 七、总结和下一步打算 ====================
add_title(doc, '七、总结和下一步打算', 1)
doc.add_paragraph()

add_title(doc, '1. 总结', 2)
doc.add_paragraph()

add_para(doc, '通过本次QC活动，成功开发了"贵阳南检修信息管理平台"，实现了以下成果：')

add_para(doc, '（1）建立了集中存储、标准统一的数据中心，实现了车间管理数据的集中存储和共享利用。')

add_para(doc, '（2）开发了20个功能模块，涵盖车间管理、生产系统、规章制度三大类，满足了车间日常管理和生产的需要。')

add_para(doc, '（3）实现了PC端和移动端的多终端访问，方便了现场作业人员的使用。')

add_para(doc, '（4）显著提升了工作效率，减少了重复填报工作，降低了职工工作负担。')

add_para(doc, '（5）通过活动开展，提高了小组成员的软件开发能力和QC活动开展水平。')

add_para(doc, '项目统计数据：')

add_para(doc, '• 开发代码量：约18000行Python代码、426个HTML模板文件')

add_para(doc, '• 数据模型：50+个数据表，涵盖人员、设备、生产等各类数据')

add_para(doc, '• 功能模块：20个功能模块')

add_para(doc, '• 开发周期：24个月')

doc.add_paragraph()
add_title(doc, '2. 下一步打算', 2)
doc.add_paragraph()

add_para(doc, '（1）持续优化：继续跟进用户反馈内容、意见，对"管理平台"各模块进行优化，根据实际情况对模块进行增减，确保符合车间使用要求。')

add_para(doc, '（2）推广应用：在段级不同车间进行推广、宣传、试用，进一步扩大"管理平台"的影响范围。')

add_para(doc, '（3）功能扩展：根据用户需求，持续增加新的功能模块，如数据分析、智能预警等。')

add_para(doc, '（4）系统升级：考虑采用更先进的技术架构，提升系统的性能和安全性。')

add_para(doc, '（5）标准建设：将成功经验形成技术标准和管理制度，为其他单位提供参考。')

doc.add_paragraph()
add_para(doc, '【建议插入图片：系统主界面截图】')

# 保存文档
output_path = r"D:\桌面\QC\贵阳南检修信息管理平台-QC成果报告.docx"
doc.save(output_path)
print(f"文档已保存到: {output_path}")